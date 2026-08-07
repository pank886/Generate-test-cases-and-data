"""文件文本提取（PDF/MD/TXT/DOCX）+ doc_id 生成。"""

import os

from observability import get_logger

logger = get_logger(__name__)


def _extract_text(file_path: str) -> str:
    """通用文本提取（支持 PDF/MD/TXT/DOCX）。"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        texts = []
        for p in reader.pages:
            t = p.extract_text()
            if t: texts.append(t)
        return "\n\n".join(texts)
    elif ext in (".md", ".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def _extract_docx(file_path: str) -> str:
    """提取 Word 文档文本，附带图片占位标记（供后续多模态替换）。"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)

    parts = []
    img_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("[表格]\n" + "\n".join(rows))

    # 提取图片（先记录位置，内容暂用占位）
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    img_dir = _docx_img_dir(file_path)
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype):
            img_index += 1
            # 保存图片到临时目录（供后续多模态模型使用）
            img_data = rel.target_part.blob
            os.makedirs(img_dir, exist_ok=True)
            img_path = os.path.join(img_dir, f"{os.path.basename(file_path)}_{img_index}.png")
            with open(img_path, "wb") as f:
                f.write(img_data)
            parts.append(f"[图片_{img_index}: {os.path.basename(img_path)}]")

    result = "\n\n".join(parts)
    if img_index > 0:
        result += f"\n\n[本文档包含 {img_index} 张图片，已保存至 {os.path.basename(img_dir)}/ 目录]"
    return result


def _docx_img_dir(file_path: str) -> str:
    """获取 docx 图片临时目录（含文件标识，防并发冲突）。"""
    stem = os.path.splitext(os.path.basename(file_path))[0]
    safe_stem = stem.replace(" ", "_").replace(".", "_")
    return os.path.join(os.path.dirname(file_path), f"_images_{safe_stem}")


def _safe_doc_id(prefix: str, *parts: str) -> str:
    """生成唯一的 doc_id（用于删除+写入的幂等操作）。

    幂等性要求：同一文件→同一 doc_id→delete_by_doc_id 清理旧数据→写入新数据。
    TODO(多用户): doc_id 追加 user_id 或 hash 后缀防跨用户碰撞，同时保持同文件幂等。
    """
    import hashlib
    _HTML_CHARS = {'<': '_lt_', '>': '_gt_', '"': '_quot_', '&': '_amp_'}
    sanitized = []
    for p in parts:
        if not p:
            continue
        s = p.replace('/', '_').replace('\\', '_').replace('$', '_')
        for ch, repl in _HTML_CHARS.items():
            s = s.replace(ch, repl)
        sanitized.append(s)
    if not sanitized:
        return prefix
    raw = prefix + "_" + "_".join(sanitized)
    # 限制总长度 ≤ 180（String(200) 留余量给 ChromaDB 内部后缀）
    if len(raw) > 180:
        suffix = hashlib.md5(raw.encode()).hexdigest()[:8]
        logger.warning("doc_id 超长截断（%d > 180）: %s… → …_%s", len(raw), raw[:60], suffix)
        return raw[:172] + "_" + suffix
    return raw
