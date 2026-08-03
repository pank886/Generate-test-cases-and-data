"""Phase A: 智能文档处理入口（替代旧 ingest_file.py 的单 Collection 流程）

入库流程：
  文件 → LLM 提取 → 1. SQLite（文档元数据 + 绑定关系 + 术语）
                   → 2. ChromaDB（纯文本 + 向量，不含业务关系）

使用方式：
    from ingest_v2 import process_product_doc, process_api_doc
    process_product_doc("uploads/doc.pdf")
    process_api_doc("uploads/api.md")
"""

import os
import re

from langchain_text_splitters import RecursiveCharacterTextSplitter

from observability import get_logger
from agent_components.dual_chroma import get_chroma_db
from agent_components.nodes import ChatTestAgentGraph
from prompts.response_model import DocModuleExtract, ApiDefExtract
from prompts.extraction_prompts import (
    product_doc_extract_prompt,
    api_def_extract_prompt,
)
import config

logger = get_logger(__name__)


def _merge_api_defs(existing: dict, incoming: dict) -> dict:
    """合并同一接口的两个版本（method+url 相同），而非简单覆盖。

    合并策略：
      - parameters/returns: 取两套字段的并集，incoming 的字段优先
      - description: 取更详细（更长）的那一个
      - name/method/url: 保留 incoming（新版本为准）
    """
    merged = dict(incoming)  # 以新版本为基底
    # parameters: 新版为 list，旧版为 dict；list 直接覆盖，dict 做字段级合并
    incoming_params = incoming.get("parameters", []) or []
    if isinstance(incoming_params, list):
        merged["parameters"] = incoming_params
    else:
        existing_params = existing.get("parameters", {}) or {}
        if isinstance(existing_params, dict):
            merged_params = dict(existing_params)
            merged_params.update(incoming_params or {})
            merged["parameters"] = merged_params
        else:
            merged["parameters"] = incoming_params

    # returns 同理
    incoming_returns = incoming.get("returns", []) or []
    if isinstance(incoming_returns, list):
        merged["returns"] = incoming_returns
    else:
        existing_returns = existing.get("returns", {}) or {}
        if isinstance(existing_returns, dict):
            merged_returns = dict(existing_returns)
            merged_returns.update(incoming_returns or {})
            merged["returns"] = merged_returns
        else:
            merged["returns"] = incoming_returns

    # description 保留更详细的那个
    desc_existing = (existing.get("description") or "").strip()
    desc_incoming = (incoming.get("description") or "").strip()
    merged["description"] = desc_incoming if len(desc_incoming) >= len(desc_existing) else desc_existing

    return merged


def _extract_text(file_path: str) -> str:
    """通用文本提取（支持 PDF/MD/TXT/DOCX）。"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        texts = []
        for p in reader.pages:
            t = p.extract_text()
            if t: texts.append(t)
        return "\n\n".join(texts)
    elif ext in (".md", ".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {ext}")


def _extract_docx(file_path: str) -> str:
    """提取 Word 文档文本，附带图片占位标记（供后续多模态替换）。"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)

    parts = []
    img_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("[表格]\n" + "\n".join(rows))

    # 提取图片（先记录位置，内容暂用占位）
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    img_dir = _docx_img_dir(file_path)
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype):
            img_index += 1
            # 保存图片到临时目录（供后续多模态模型使用）
            img_data = rel.target_part.blob
            os.makedirs(img_dir, exist_ok=True)
            img_path = os.path.join(img_dir, f"{os.path.basename(file_path)}_{img_index}.png")
            with open(img_path, "wb") as f:
                f.write(img_data)
            parts.append(f"[图片_{img_index}: {os.path.basename(img_path)}]")

    result = "\n\n".join(parts)
    if img_index > 0:
        result += f"\n\n[本文档包含 {img_index} 张图片，已保存至 {os.path.basename(img_dir)}/ 目录]"
    return result


def _docx_img_dir(file_path: str) -> str:
    """获取 docx 图片临时目录（含文件标识，防并发冲突）。"""
    stem = os.path.splitext(os.path.basename(file_path))[0]
    safe_stem = stem.replace(" ", "_").replace(".", "_")
    return os.path.join(os.path.dirname(file_path), f"_images_{safe_stem}")


def _safe_doc_id(prefix: str, *parts: str) -> str:
    """生成唯一的 doc_id（用于删除+写入的幂等操作）。

    幂等性要求：同一文件→同一 doc_id→delete_by_doc_id 清理旧数据→写入新数据。
    TODO(多用户): doc_id 追加 user_id 或 hash 后缀防跨用户碰撞，同时保持同文件幂等。
    """
    import hashlib
    _HTML_CHARS = {'<': '_lt_', '>': '_gt_', '"': '_quot_', '&': '_amp_'}
    sanitized = []
    for p in parts:
        if not p:
            continue
        s = p.replace('/', '_').replace('\\', '_').replace('$', '_')
        for ch, repl in _HTML_CHARS.items():
            s = s.replace(ch, repl)
        sanitized.append(s)
    if not sanitized:
        return prefix
    raw = prefix + "_" + "_".join(sanitized)
    # 限制总长度 ≤ 180（String(200) 留余量给 ChromaDB 内部后缀）
    if len(raw) > 180:
        suffix = hashlib.md5(raw.encode()).hexdigest()[:8]
        logger.warning("doc_id 超长截断（%d > 180）: %s… → …_%s", len(raw), raw[:60], suffix)
        return raw[:172] + "_" + suffix
    return raw


def _extract_valid_api_paths(full_text: str) -> set[tuple[str, str]]:
    """从 yapi 导出的 MD 文档中提取所有合法的 (METHOD, URL) 白名单。

    以 ``**Path：** `` 为锚点，向后搜索最近 500 字符内的 ``**Method：** ``。
    返回 {(METHOD_UPPER, url), ...} 集合，用于过滤 LLM 幻觉的接口。

    设计意图：LLM 可能因参数字段名（deviceId）或记忆污染（跨项目接口）
    编造不存在的接口。白名单直接扫描原始文档的 Path/Method 行，
    不依赖 LLM 质量，提供一道独立防线。
    """
    path_re = re.compile(r'\*\*Path[：:]\*\*\s+(/\S+)')
    method_re = re.compile(r'\*\*Method[：:]\*\*\s+(\w+)')

    valid: set[tuple[str, str]] = set()
    for path_m in path_re.finditer(full_text):
        url = path_m.group(1).strip().rstrip("/")
        if not url:
            continue
        pos = path_m.end()
        # 向后搜索最近 500 字符内的 Method（yapi 格式 Path 在前 Method 在后）
        search_end = min(len(full_text), pos + 500)
        search_text = full_text[pos:search_end]
        method_m = method_re.search(search_text)
        if method_m:
            method = method_m.group(1).strip().upper()
            valid.add((method, url))
    return valid


def _cascade_bind_to_module_docs(session, doc_type: str, doc_id: str, module_name: str):
    """级联关联：文档绑定模块时，自动与该模块下所有异类文档建立 doc↔doc 绑定。"""
    from database.operations import BindingOps
    bound_docs = BindingOps.get_bound_docs(session, module_name)
    for other_doc in bound_docs:
        if other_doc.doc_type != doc_type and other_doc.id != doc_id:
            BindingOps.bind(session, doc_type, doc_id, other_doc.doc_type, other_doc.id)


def _delete_sqlite_doc(doc_id: str):
    """删除 SQLite 中的文档记录（作为 ChromaDB 写入失败的补偿动作）。"""
    from database import get_session_ctx
    from database.operations import DocOps
    try:
        with get_session_ctx() as session:
            DocOps.delete_document(session, doc_id)
            logger.info("   [补偿] 已回滚 SQLite 记录: %s", doc_id)
    except Exception as e:
        logger.error("   [补偿] SQLite 回滚失败（需人工清理）: %s - %s", doc_id, e, exc_info=True)


def _save_to_sqlite(doc_id: str, file_name: str, file_type: str, doc_type: str,
                    chunk_count: int, module_name: str = "",
                    glossary_terms: list = None):
    """写入 SQLite：文档记录 + 术语。

    必须在 ChromaDB 写入**之前**调用。若后续 ChromaDB 写入失败，
    由调用方通过 _delete_sqlite_doc() 执行补偿回滚。
    module_name 仅用于日志，不做自动绑定（由用户在前端手动关联）。
    """
    from database import get_session_ctx
    from database.operations import DocOps, GlossaryOps

    try:
        with get_session_ctx() as session:
            # 1. 文档记录（session.merge() 不触发 column default，显式设置 upload_time）
            from database.models import Document
            from datetime import datetime, timezone
            doc = Document(
                id=doc_id, file_name=file_name, file_type=file_type,
                doc_type=doc_type, chunk_count=chunk_count, status="pending",
                upload_time=datetime.now(timezone.utc),
            )
            session.merge(doc)

            # 2. 术语（如果提供了）
            if glossary_terms:
                GlossaryOps.replace_terms(
                    session, doc_id, glossary_terms,
                    source_doc=file_name,
                )

            if module_name:
                logger.debug(f"   [SQLite] 文档 {doc_id} 关联模块: {module_name}")
    except Exception:
        logger.error("   [SQLite] 写入失败", exc_info=True)
        raise


def _save_single_chunk(doc_id: str, chunk_index: int, content: str,
                        page_name: str = ""):
    """写入单条 document_chunk 记录（供 Axure 逐页写入）。"""
    from database import get_session_ctx
    from database.models import DocumentChunk
    try:
        with get_session_ctx() as session:
            chunk = DocumentChunk(
                doc_id=doc_id,
                chunk_index=chunk_index,
                content=content,
                page_name=page_name,
                token_count=len(content),
            )
            session.add(chunk)
            session.commit()
    except Exception:
        logger.error("   [document_chunks] 单条写入失败: %s[%d]", doc_id, chunk_index, exc_info=True)
        raise


def _save_document_chunks(doc_id: str, chunks: list[str], page_name: str = ""):
    """写入 document_chunks 表（原文 + 摘要占位）。"""
    from database import get_session_ctx
    from database.models import DocumentChunk
    try:
        with get_session_ctx() as session:
            for i, content in enumerate(chunks):
                chunk = DocumentChunk(
                    doc_id=doc_id,
                    chunk_index=i,
                    content=content,
                    page_name=page_name,
                    token_count=len(content),
                )
                session.add(chunk)
            session.commit()
        logger.info(f"   [document_chunks] 写入 {len(chunks)} 条")
    except Exception:
        logger.error("   [document_chunks] 写入失败", exc_info=True)
        raise


def _delete_document_chunks(doc_id: str):
    """删除 document_chunks 记录（ChromaDB 写入失败时的补偿动作）。"""
    from database import get_session_ctx
    from database.models import DocumentChunk
    try:
        with get_session_ctx() as session:
            session.query(DocumentChunk).filter_by(doc_id=doc_id).delete()
            session.commit()
        logger.info("   [补偿] 已回滚 document_chunks: %s", doc_id)
    except Exception as e:
        logger.error("   [补偿] document_chunks 回滚失败: %s - %s", doc_id, e, exc_info=True)


def _parse_chunk_summaries(text: str, expected_count: int) -> list[str]:
    """正则解析 ===CHUNK_SUMMARY=== 分隔的摘要文本。

    返回长度等于 expected_count 的列表，不足补空字符串。
    """
    import re
    # 按 ===CHUNK_SUMMARY=== 拆分，取每段的第一行
    parts = re.split(r'===CHUNK_SUMMARY===', text)
    summaries = [p.strip().split('\n')[0].strip() for p in parts[1:]]  # 第一个是空
    # 补齐
    while len(summaries) < expected_count:
        summaries.append("")
    return summaries[:expected_count]


def _generate_batch_summaries(doc_id: str, chunks: list[str], file_name: str,
                               progress_cb=None, page_name: str = ""):
    """批量生成 simple_summary（5 chunks/批，同步等待，失败写补偿任务）。

    LLM 输出 ===CHUNK_SUMMARY=== 分隔词，正则匹配解析。
    """
    import re
    import config
    from agent_components.nodes import _get_llm
    from prompts.extraction_prompts import batch_chunk_summary_prompt
    from database import get_session_ctx
    from database.models import DocumentChunk
    from database.operations.compensation import CompensationOps

    cb = progress_cb or (lambda p, m: None)
    total = len(chunks)
    batch_size = config.BATCH_SUMMARY_CHUNK_SIZE
    prompt = batch_chunk_summary_prompt()
    llm = _get_llm()

    failed_indices = []
    batch_count = (total + batch_size - 1) // batch_size

    for bi in range(0, total, batch_size):
        batch_end = min(bi + batch_size, total)
        batch_chunks = chunks[bi:batch_end]
        batch_num = bi // batch_size + 1
        pct = 70 + int((batch_num / batch_count) * 15)
        cb(pct, f"AI 生成摘要 ({batch_num}/{batch_count})...")

        chunks_text = "\n\n".join(
            f"[块{bi+j}] {c}" for j, c in enumerate(batch_chunks)
        )
        try:
            result = llm.invoke(prompt.format_messages(
                file_name=file_name,
                start_idx=bi,
                end_idx=batch_end - 1,
                total=total,
                page_name=page_name,
                chunks=chunks_text,
            ))
            llm_text = result.content if hasattr(result, "content") else str(result)
            summaries = _parse_chunk_summaries(llm_text, len(batch_chunks))
        except Exception as e:
            logger.warning("   [摘要] 批次 %d/%d LLM 调用失败: %s", batch_num, batch_count, e)
            # 整批失败 → 全部标记补偿
            summaries = [""] * len(batch_chunks)
            failed_indices.extend(range(bi, batch_end))

        # 写入 document_chunks.simple_summary
        try:
            with get_session_ctx() as session:
                for j, summary in enumerate(summaries):
                    if summary:  # 非空才写入
                        session.query(DocumentChunk).filter_by(
                            doc_id=doc_id, chunk_index=bi + j,
                        ).update({"simple_summary": summary})
                session.commit()
        except Exception as e:
            logger.error("   [摘要] 写入 SQLite 失败: %s", e)

        # 收集解析失败的索引（空摘要）
        for j, s in enumerate(summaries):
            if not s and (bi + j) not in failed_indices:
                failed_indices.append(bi + j)

    # 创建补偿任务（失败的 chunk）
    if failed_indices:
        try:
            with get_session_ctx() as session:
                CompensationOps.create(
                    session, "simple_summary",
                    {"doc_id": doc_id, "chunk_indices": failed_indices,
                     "file_name": file_name},
                    max_retries=config.COMPENSATION_MAX_RETRIES,
                )
                session.commit()
            logger.info("   [补偿] 已创建 simple_summary 补偿任务: %d 个 chunk", len(failed_indices))
        except Exception as e:
            logger.error("   [补偿] 创建补偿任务失败: %s", e)


def process_product_doc(file_path: str, progress_cb=None) -> dict:
    """处理产品文档：提取文本 -> LLM 提取模块关联 -> 存入 product_docs + SQLite。

    Args:
        progress_cb: 可选，进度回调 (0~100, message)
    """
    cb = progress_cb or (lambda p, m: None)
    from observability import log_phase_header
    log_phase_header("Phase A — 文档摄入与向量化")
    logger.info(f"\n{'=' * 60}")
    logger.info(f"[Phase A] 处理产品文档: {os.path.basename(file_path)}")

    db = get_chroma_db()
    graph = ChatTestAgentGraph()
    file_name = os.path.basename(file_path)
    file_type = os.path.splitext(file_path)[1].lstrip(".")

    # 1. 提取文本
    cb(5, "提取文本中...")
    full_text = _extract_text(file_path).strip()
    if not full_text:
        raise ValueError("文档内容为空")
    logger.info(f"   => 提取文本 {len(full_text)} 字符")

    # _extract_text 可能产生临时图片目录，统一在 finally 中清理
    _img_dir = _docx_img_dir(file_path) if os.path.splitext(file_path)[1].lower() == ".docx" else None
    try:
        # 2. 切块（前置，后续 LLM 提取和 ChromaDB 入库共用同一批块）
        cb(15, "文本切分中...")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", "。", "，", " "],
        )
        chunks = splitter.split_text(full_text)
        logger.info(f"   => 切分为 {len(chunks)} 个文本块")

        # 将块打包为不超过 MAX_INGEST_CHARS_PER_BATCH 的批次
        def _group_chunks_into_batches(chunks: list[str], max_chars: int) -> list[str]:
            """分组拼接为每批不超过 max_chars 的文本段。"""
            out, batch = [], []
            for c in chunks:
                candidate = "\n\n".join(batch + [c]) if batch else c
                if len(candidate) > max_chars and batch:
                    out.append("\n\n".join(batch))
                    batch = [c]
                else:
                    batch.append(c)
            if batch:
                out.append("\n\n".join(batch))
            return out

        batch_limit = config.MAX_INGEST_CHARS_PER_BATCH
        text_batches = _group_chunks_into_batches(chunks, batch_limit) if len(full_text) > batch_limit else [full_text]
        logger.info(f"   => 打包为 {len(text_batches)} 批（每批 ≤ {batch_limit} 字符）")

        # 3. LLM 提取模块信息（分批处理，合并 related_modules / tags）
        cb(30, "AI 分析模块信息...")
        prompt = product_doc_extract_prompt()
        logger.info("   => LLM 提取模块信息...")
        module_name = ""
        related: set[str] = set()
        business_summary = ""
        tags: set[str] = set()
        for bi, batch_text in enumerate(text_batches, 1):
            result = graph._invoke_structured(
                prompt, DocModuleExtract,
                method="json_mode",
                doc_text=batch_text,
            )
            if not module_name and result.module_name:
                module_name = result.module_name
            if result.related_modules:
                related.update(result.related_modules)
            if not business_summary and result.business_summary:
                business_summary = result.business_summary
            if result.tags:
                tags.update(result.tags)
            if len(text_batches) > 1:
                logger.info(f"   [{bi}/{len(text_batches)}] 模块: {result.module_name or '?'}, "
                            f"+{len(result.related_modules or [])} 关联")
        module_name = module_name or "Unknown"
        related_list = sorted(related)
        logger.info(f"   => 模块: {module_name}, 关联: {related_list}, 标签: {sorted(tags)}")

        # 4. LLM 提取业务术语表（分批处理，合并去重）
        cb(50, "AI 提取术语表...")
        from prompts.response_model import GlossaryExtract
        from prompts.extraction_prompts import glossary_extract_prompt
        terms = []
        try:
            glossary_prompt = glossary_extract_prompt()
            seen_terms: set[str] = set()
            for bi, batch_text in enumerate(text_batches, 1):
                glossary_result = graph._invoke_structured(
                    glossary_prompt, GlossaryExtract,
                    method="json_mode",
                    doc_text=batch_text,
                )
                batch_terms = glossary_result.terms if hasattr(glossary_result, "terms") else []
                for t in batch_terms:
                    key = (t.get("term", t.get("name", "")).strip())
                    if key and key not in seen_terms:
                        seen_terms.add(key)
                        terms.append(t)
                if len(text_batches) > 1:
                    logger.info(f"   [{bi}/{len(text_batches)}] 术语: +{len(batch_terms)} 条（合并后 {len(terms)} 条）")
            if terms:
                logger.info(f"   => 术语表: {len(terms)} 条")
        except Exception as e:
            logger.warning("术语表提取跳过: %s", e, exc_info=True)

        # 5. 写入 SQLite document_chunks 表（原文 + 摘要占位）
        cb(60, "写入文档块...")
        doc_id = _safe_doc_id("prod", file_name, module_name)
        _save_document_chunks(doc_id, chunks, page_name="")
        # 写 documents 表（元数据）
        _save_to_sqlite(
            doc_id=doc_id,
            file_name=file_name,
            file_type=file_type,
            doc_type="product",
            chunk_count=len(chunks),
            module_name=module_name,
            glossary_terms=terms,
        )
        logger.info(f"   [SQLite] document_chunks + documents 入库完成 (doc_id={doc_id})")

        # 6. 批量生成 simple_summary（同步等待，5 chunks/批）
        cb(70, "AI 生成摘要...")
        _generate_batch_summaries(doc_id, chunks, file_name, progress_cb=cb)

        # 7. 写入 ChromaDB（检索文本，失败时补偿回滚 SQLite）
        cb(90, "向量化入库中...")
        try:
            from database import get_session_ctx
            from database.models import DocumentChunk
            # 从 document_chunks 读摘要构造检索文本
            chunk_records = []
            with get_session_ctx() as session:
                chunk_records = session.query(DocumentChunk).filter_by(
                    doc_id=doc_id).order_by(DocumentChunk.chunk_index).all()
                # detach from session
                chunk_records = [{
                    "content": c.content,
                    "simple_summary": c.simple_summary,
                    "page_name": c.page_name,
                } for c in chunk_records]

            search_texts = [_build_doc_search_text(c) for c in chunk_records]
            db.delete_by_doc_id(doc_id)
            db.add_product_doc_chunks(doc_id, search_texts)
            logger.info(f"   [ChromaDB] 入库完成 (doc_id={doc_id}, 检索文本)")
        except Exception:
            logger.error("   [ChromaDB] 写入失败，启动补偿回滚 SQLite", exc_info=True)
            _delete_sqlite_doc(doc_id)
            _delete_document_chunks(doc_id)
            raise

        cb(95, "入库完成")
        return {
            "doc_id": doc_id,
            "module_name": module_name,
            "related_modules": related_list,
            "chunks": len(chunks),
        }
    finally:
        if _img_dir and os.path.isdir(_img_dir):
            import shutil as _su
            _su.rmtree(_img_dir, ignore_errors=True)
            logger.debug("已清理临时图片目录: %s", _img_dir)


def extract_apis_from_yapi_md(text: str) -> list[dict]:
    """纯代码提取：解析 YApi 导出的 MD 文档，返回接口定义列表。

    适用于格式规整的 YApi MD，不需要 LLM。
    提取字段：name, url, method, description, headers, parameters, returns
    """
    import re as _re
    from bs4 import BeautifulSoup

    # ── 提取模块名（从 h1 标签取纯文本）──
    module_name = ""
    h1_match = _re.search(r'<h1[^>]*>(.+?)</h1>', text)
    if h1_match:
        module_name = _re.sub(r'<[^>]+>', '', h1_match.group(1)).strip()

    # ── 切分 API 段（跳过 h1/# 前言，不要混入第一个 API）──
    parts = _re.split(r'(?=\n## )', text)
    parts = [p for p in parts if p.strip() and p.strip().startswith('## ')]

    apis = []
    for part in parts:
        part = part.strip()
        if not part:
            continue

        # ── 基本信息 ──
        first_line = part.split('\n')[0].strip()
        name = _re.sub(r'^##\s+', '', first_line).strip()

        url_match = _re.search(r'\*\*Path：\*\*\s*(.+)', part)
        url = url_match.group(1).strip() if url_match else ""

        method_match = _re.search(r'\*\*Method：\*\*\s*(.+)', part)
        method = method_match.group(1).strip().upper() if method_match else "?"

        # 接口描述：YApi 空描述导出为 <p></p>，正则捕获后再剥掉 HTML 标签残留（不能直接留 </p>）
        desc_match = _re.search(r'\*\*接口描述：\*\*\s*\n?\s*(?:<p[^>]*>)?(.*?)(?:</p>)?\s*\n', part)
        description = ""
        if desc_match:
            description = _re.sub(r'<[^>]+>', '', desc_match.group(1)).strip()

        # ── 辅助：解析 HTML 表格为参数数组 ──
        def _parse_html_table(html_str: str) -> list[dict]:
            """解析 HTML <table> 为 [{name, type, required, description, default, children}]"""
            soup = BeautifulSoup(html_str, 'html.parser')
            table = soup.find('table')
            if not table:
                return []
            # 找表头确定列映射
            headers = []
            for th in table.find_all('th'):
                key = th.get('key', th.get_text(strip=True))
                headers.append(key)
            if not headers:
                return []
            # 列名映射（YApi 格式：名称/类型/是否必须/默认值/备注）
            col_map = {'name': -1, 'type': -1, 'required': -1, 'default': -1, 'desc': -1}
            for idx, h in enumerate(headers):
                hl = h.lower()
                if '名称' in h or 'name' in hl or '字段' in h or '参数' in h:
                    col_map['name'] = idx
                elif '类型' in h or 'type' in hl:
                    col_map['type'] = idx
                elif '必须' in h or 'required' in hl:
                    col_map['required'] = idx
                elif '默认' in h or 'default' in hl:
                    col_map['default'] = idx
                elif '备注' in h or 'desc' in hl or '说明' in h:
                    col_map['desc'] = idx

            # 解析行，通过缩进判断层级
            rows = table.find_all('tr')
            result = []
            stack = [(result, -1)]  # (parent_list, indent_level)

            for tr in rows:
                cells = tr.find_all('td')
                if len(cells) < 2:
                    continue
                # 计算缩进层级
                first_cell = cells[0]
                spans = first_cell.find_all('span')
                indent = 0
                for sp in spans:
                    style = sp.get('style', '')
                    if 'padding-left' in style:
                        try:
                            px = int(_re.search(r'padding-left:\s*(\d+)px', style).group(1))
                            indent = max(indent, px // 20)  # 每 20px 一级
                        except (ValueError, AttributeError):
                            pass

                # 提取字段值
                def _cell_text(idx):
                    if idx < 0 or idx >= len(cells):
                        return ''
                    # 去掉嵌套的 span 样式文本，只取直接文本或 API 名称
                    t = cells[idx].get_text(separator=' ', strip=True)
                    # 清理树形连接符
                    t = _re.sub(r'^\s*[├└]─?\s*', '', t)
                    return t.strip()

                name_val = _cell_text(col_map['name'])
                type_val = _cell_text(col_map['type'])
                required_str = _cell_text(col_map['required'])
                default_val = _cell_text(col_map['default'])
                desc_val = _cell_text(col_map['desc'])

                if not name_val:
                    continue

                required = '必须' in required_str or required_str.lower() == '是' or required_str.lower() == 'true'

                item = {
                    'name': name_val,
                    'type': type_val or 'string',
                    'required': required,
                    'description': desc_val,
                    'default': default_val or None,
                }

                # 处理层级：弹出比当前缩进更深的栈
                while len(stack) > 1 and stack[-1][1] >= indent:
                    stack.pop()
                parent_list = stack[-1][0]
                parent_list.append(item)
                # 如果有嵌套类型，准备 children
                if type_val and ('object' in type_val.lower() or '[]' in type_val or 'array' in type_val.lower()):
                    item['children'] = []
                    stack.append((item['children'], indent))

            return result

        # ── 解析请求参数（支持 Headers/Query/Body；Markdown 表格或 HTML 表格）──
        headers_list = []
        params_list = []

        # 辅助：解析 YApi 导出的 Markdown 参数表（按表头自动映射列名）
        def _parse_md_table(md_str: str) -> list[dict]:
            lines = [ln for ln in md_str.strip().split('\n') if ln.strip().startswith('|')]
            if len(lines) < 2:
                return []
            rows = [[c.strip() for c in ln.strip().strip('|').split('|')] for ln in lines]
            header = rows[0]
            # 列映射：只映射第一个命中，避免「参数名称/参数值」都含"参数"时冲突
            col_map = {'name': -1, 'type': -1, 'required': -1, 'default': -1, 'desc': -1}
            for idx, h in enumerate(header):
                hl = h.lower()
                if col_map['name'] == -1 and ('名称' in h or 'name' in hl or '参数' in h or '字段' in h):
                    col_map['name'] = idx
                elif col_map['type'] == -1 and ('类型' in h or 'type' in hl):
                    col_map['type'] = idx
                elif col_map['required'] == -1 and ('必须' in h or 'required' in hl):
                    col_map['required'] = idx
                elif col_map['default'] == -1 and ('默认' in h or 'default' in hl):
                    col_map['default'] = idx
                elif col_map['desc'] == -1 and ('备注' in h or '说明' in h or 'desc' in hl or '描述' in h):
                    col_map['desc'] = idx
            result = []
            for row in rows[1:]:
                if all(set(c) <= {'-', ':', ' '} for c in row if c):
                    continue  # 分隔行
                def _cell(idx):
                    return row[idx] if 0 <= idx < len(row) else ''
                name_val = _cell(col_map['name'])
                if not name_val:
                    continue
                req_str = _cell(col_map['required'])
                result.append({
                    'name': name_val,
                    'type': _cell(col_map['type']) or 'string',
                    'required': '必须' in req_str or req_str.lower() == '是' or req_str.lower() == 'true',
                    'description': _cell(col_map['desc']),
                    'default': _cell(col_map['default']) or None,
                })
            return result

        # 请求参数区（### 请求参数 → ### 返回数据），只在区内解析，避免误取返回数据表
        req_section = ''
        req_match = _re.search(r'### 请求参数\s*\n(.*?)(?=\n### 返回数据|\Z)', part, _re.DOTALL)
        if req_match:
            req_section = req_match.group(1)

        def _subsection(sec: str, title: str) -> str:
            """取 **title** 标题后的子段内容（到下一个 **标题 或 ### 或结尾）。"""
            m = _re.search(r'\*\*' + _re.escape(title) + r'\*\*\s*\n(.*?)(?=\n\*\*|\n###\s|\Z)',
                           sec, _re.DOTALL)
            return m.group(1) if m else ''

        hdr_section = _subsection(req_section, 'Headers')
        query_section = _subsection(req_section, 'Query')
        body_section = _subsection(req_section, 'Body')

        headers_list = _parse_md_table(hdr_section) if hdr_section else []
        if query_section:
            params_list.extend(_parse_md_table(query_section))
        if body_section:
            body_params = _parse_md_table(body_section)
            params_list.extend(body_params or _parse_html_table(body_section))

        # ── 解析返回数据 ──
        returns_list = []
        ret_match = _re.search(r'### 返回数据\s*\n(.*?)(?=\n##|\Z)', part, _re.DOTALL)
        if ret_match:
            returns_html = ret_match.group(1)
            returns_list = _parse_html_table(returns_html)

        api = {
            'name': name,
            'url': url,
            'method': method,
            'description': description or name,
            'headers': headers_list,
            'parameters': params_list,
            'returns': returns_list,
            'annotations': {},
        }
        apis.append(api)

    return {"apis": apis, "module_name": module_name}


def _split_text_by_headers(text: str, max_chars: int) -> list:
    """按 ## 标题切分文本，每个 API 独立成段。

    只按 ## (h2) 切——保证每个 API 完整不被截断。
    不拼批次：每个 API 就是一段，不再按字符数合并。
    max_chars 仅用于单个 API 超出限制时的截断保护。
    """
    import re
    parts = re.split(r'(?=\n## )', text)
    # 第一段可能是 # 标题行，不含 API；合并到第一个 ## 段
    if len(parts) >= 2:
        pre = parts[0].strip()
        if pre and not pre.startswith('## '):
            parts[1] = pre + "\n\n" + parts[1].lstrip()
        parts = parts[1:] if not parts[0].strip().startswith('## ') else parts

    batches = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        batches.append(part)
    return batches


def process_api_doc(file_path: str, default_module: str = None, progress_cb=None) -> dict:
    """处理接口文档：提取 → 入库（委托给新版函数）。

    .. deprecated::
       使用 ``process_api_doc_extract()`` + ``commit_api_docs()`` 替代。
       当前保留仅用于 CLI 兼容，内部已委托新版函数。
    """
    logger.warning("process_api_doc() 已弃用，请使用 process_api_doc_extract() + commit_api_docs()")
    cb = progress_cb or (lambda p, m: None)

    # 委托新版提取
    cb(5, "提取接口定义...")
    extracted = process_api_doc_extract(file_path, default_module=default_module)
    apis = extracted.get("apis", [])
    module = extracted.get("module_name") or default_module or "Unknown"
    if not apis:
        logger.warning("未提取到接口定义")
        return {"doc_id": "", "module_name": module, "api_count": 0}

    # 委托新版入库（已含 SQLite 先写 + ChromaDB 补偿逻辑）
    cb(80, "入库中...")
    result = commit_api_docs(file_path, module, apis)
    logger.info("   => 委托入库完成: %d 个接口", result["api_count"])
    return result


def process_api_doc_extract(file_path: str, default_module: str = None,
                             progress_cb=None) -> dict:
    """Phase 1: 提取接口列表（不入库），返回给前端确认。

    Returns: {"module_name": str, "apis": [dict], "file_name": str}
    """
    cb = progress_cb or (lambda p, m: None)
    logger.info(f"[Phase A] 提取接口: {os.path.basename(file_path)}")

    graph = ChatTestAgentGraph()
    file_name = os.path.basename(file_path)

    cb(5, "读取文档...")
    full_text = _extract_text(file_path).strip()
    if not full_text:
        raise ValueError("文档内容为空")

    batch_limit = config.MAX_INGEST_CHARS_PER_BATCH
    batches = _split_text_by_headers(full_text, batch_limit) if len(full_text) > batch_limit else [full_text]

    all_apis = []
    module = default_module
    total = len(batches)
    done = [0]  # 列表包装，实现跨线程计数

    import concurrent.futures
    from threading import Lock
    _lock = Lock()

    def _extract_one(batch_text: str) -> tuple:
        """单个 API 提取（在线程池中并发执行）。"""
        prompt = api_def_extract_prompt()
        result = graph._invoke_structured(
            prompt, ApiDefExtract, method="json_mode", doc_text=batch_text,
        )
        mod = result.module_name
        apis_raw = result.apis if hasattr(result, "apis") else []
        apis = [a.model_dump() if hasattr(a, "model_dump") else a for a in apis_raw]
        with _lock:
            done[0] += 1
            pct = int(10 + (done[0] / total) * 70)
            cb(pct, f"AI 提取接口定义 ({done[0]}/{total})...")
        return mod, apis

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        futures = [executor.submit(_extract_one, b) for b in batches]
        for f in concurrent.futures.as_completed(futures):
            mod, apis = f.result()
            if not module:
                module = mod
            all_apis.extend(apis)

    # ---- 白名单校验：过滤 LLM 幻觉的接口 ----
    # 扫描原始文档，提取所有 **Path：** + **Method：** 对作为白名单。
    # LLM 可能因参数字段名（如 deviceStatus→编造 /device/info）或记忆污染
    # （跨项目接口混淆）而输出不存在的接口，白名单提供独立防线。
    valid_paths = _extract_valid_api_paths(full_text)
    if valid_paths:
        before = len(all_apis)
        hallucinated = []
        filtered_apis = []
        for a in all_apis:
            key = (a.get("method", "").strip().upper(), a.get("url", "").rstrip("/"))
            if key in valid_paths:
                filtered_apis.append(a)
            else:
                hallucinated.append(key)
        all_apis = filtered_apis
        removed = before - len(all_apis)
        if removed:
            logger.warning(
                "⚠️ 白名单校验: 过滤 %d 个文档中不存在的接口: %s",
                removed, [f"{m} {u}" for m, u in hallucinated[:15]],
            )
    # ---- 白名单校验结束 ----

    # 合并去重：method+url 相同时合并参数和返回值，而非简单覆盖
    merged = {}
    dup_count = 0
    for api in all_apis:
        key = f"{api.get('method', '')} {api.get('url', '')}"
        if key in merged:
            dup_count += 1
            merged[key] = _merge_api_defs(merged[key], api)
        else:
            merged[key] = api
    if dup_count:
        logger.warning("检测到 %d 个重复接口（method+url 相同），已合并参数/返回值/描述", dup_count)
    apis = list(merged.values())

    module = module or "Unknown"
    return {"module_name": module, "apis": apis, "file_name": file_name}


def _build_api_search_text(api: dict) -> str:
    """构造 API 自然语言检索文本（替代 JSON 原文写入 ChromaDB）。

    格式: "{method} {url} {api_name}。{description}。
            参数: {param_name}{param_type}{'必填' if required else ''}{param_desc}; ...
            返回值: {ret_name}{ret_type}; ..."
    """
    name = api.get("name", "")
    url = api.get("url", "")
    method = api.get("method", "?").upper()
    desc = api.get("description", name)

    parts = [f"{method} {url} {name}。{desc}。"]

    # 参数
    params = api.get("parameters", [])
    if isinstance(params, list) and params:
        param_strs = []
        for p in params[:20]:  # 最多 20 个参数，防过长
            pn = p.get("name", "")
            pt = p.get("type", "string")
            pr = "必填" if p.get("required") else ""
            pd = p.get("description", "")
            param_strs.append(f"{pn}{pt}{pr}{pd}")
        parts.append("参数: " + "; ".join(param_strs))

    # 返回值
    returns = api.get("returns", [])
    if isinstance(returns, list) and returns:
        ret_strs = [f"{r.get('name', '')}{r.get('type', '')}" for r in returns[:10]]
        parts.append("返回值: " + "; ".join(ret_strs))

    # 标签
    tags = []
    annotations = api.get("annotations", {})
    if isinstance(annotations, dict):
        if annotations.get("is_export", {}).get("active"):
            tags.append("导出接口")
        if annotations.get("has_path_params", {}).get("active"):
            tags.append("RESTful路径参数")
    if tags:
        parts.append("标签: " + " ".join(tags))

    return "。".join(parts)


def _build_doc_search_text(chunk: dict) -> str:
    """构造产品/Axure 文档检索文本。

    优先级: analyzed_summary > simple_summary > content 前 500 字
    """
    page = chunk.get("page_name", "")
    summary = chunk.get("analyzed_summary") or chunk.get("simple_summary") or ""
    if summary:
        return f"[{page}] {summary}" if page else summary
    # fallback: content 截断
    content = chunk.get("content", "")
    truncated = content[:500] if len(content) > 500 else content
    return f"[{page}] {truncated}" if page else truncated


def commit_api_docs(file_path: str, module_name: str, apis: list[dict],
                    progress_cb=None, delete_original: bool = False) -> dict:
    """Phase 2: 用户确认后，接口批量入库。

    所有 API 先批量写入 SQLite（同一事务，含 api_* 结构化列），再逐条写入 ChromaDB。
    ChromaDB 写入检索文本（自然语言），不再存原始 JSON。
    ChromaDB 任一条失败时补偿回滚所有 SQLite 记录。
    仅 delete_original=True 时删除原文件。
    """
    import json as _json
    cb = progress_cb or (lambda p, m: None)
    logger.info(f"[Phase A] 入库 {len(apis)} 个接口文档")

    db = get_chroma_db()
    file_name = os.path.basename(file_path)
    file_type = os.path.splitext(file_path)[1].lstrip(".")
    doc_ids = []

    # ---- Phase 1: 批量写入 SQLite（同一事务，含 api_* 结构化列）----
    from database import get_session_ctx
    from database.models import Document
    from datetime import datetime, timezone

    cb(10, "写入业务数据...")
    docs_to_insert = []
    for i, api in enumerate(apis):
        api_name = api.get("name", f"api_{i}")
        url = api.get("url", "")
        method = api.get("method", "?")
        doc_id = _safe_doc_id("api", file_name, module_name, method, url, api_name)
        # ★ 自动标注 API 异常标识（is_export / has_path_params 等）
        from agent_components.api_annotations import ApiAnnotationRegistry
        ApiAnnotationRegistry.apply_all(api)
        doc_ids.append(doc_id)
        docs_to_insert.append(Document(
            id=doc_id,
            file_name=f"{api.get('method', '?')} {api.get('url', '')}",
            file_type=file_type,
            doc_type="api",
            chunk_count=1,
            status="pending",
            upload_time=datetime.now(timezone.utc),
            # ── API 结构化列 ──
            api_name=api_name,
            api_url=url,
            api_method=method.upper() if method else "?",
            api_description=api.get("description", api_name),
            api_headers=_json.dumps(api.get("headers", []), ensure_ascii=False),
            api_parameters=_json.dumps(api.get("parameters", []), ensure_ascii=False),
            api_returns=_json.dumps(api.get("returns", []), ensure_ascii=False),
            api_annotations=_json.dumps(api.get("annotations", {}), ensure_ascii=False),
        ))

    try:
        with get_session_ctx() as session:
            for d in docs_to_insert:
                session.merge(d)
    except Exception:
        logger.error("   [SQLite] 批量写入失败，无数据需要补偿", exc_info=True)
        raise

    logger.info(f"   [SQLite] 批量入库完成: {len(doc_ids)} 条（含 api_* 结构化列）")

    # ---- Phase 2: 逐条写入 ChromaDB（检索文本，失败时补偿回滚 SQLite）----
    try:
        for i, api in enumerate(apis):
            api_name = api.get("name", f"api_{i}")
            url = api.get("url", "")
            method = api.get("method", "?")
            doc_id = doc_ids[i]
            pct = int(50 + (i / len(apis)) * 40)
            cb(pct, f"向量化入库 {i+1}/{len(apis)}...")

            # 构造检索文本并写入 ChromaDB
            search_text = _build_api_search_text(api)
            # 包装为兼容现有 add_api_defs 的格式：单个 dict 含检索文本
            api_for_chroma = dict(api)
            api_for_chroma["_search_text"] = search_text
            api_for_chroma["name"] = api_name
            api_for_chroma["_doc_id"] = doc_id

            db.delete_by_doc_id(doc_id)
            db.add_api_defs(doc_id, [api_for_chroma])
    except Exception:
        logger.error("   [ChromaDB] 写入失败，启动补偿回滚所有 SQLite 记录", exc_info=True)
        for did in doc_ids:
            _delete_sqlite_doc(did)
        raise

    logger.info(f"   [ChromaDB] 入库完成: {len(doc_ids)} 条（检索文本）")

    # 仅当全部接口选中时才废弃原文件
    if delete_original:
        try:
            os.remove(file_path)
            meta_path = file_path + ".meta.json"
            if os.path.exists(meta_path):
                os.remove(meta_path)
            logger.info(f"   => 已删除原文件: {file_name}")
        except OSError:
            logger.warning("原文件删除失败: %s", file_name, exc_info=True)
    else:
        logger.info(f"   => 保留原文件（部分接口未入库）: {file_name}")

    cb(95, "入库完成")
    return {"doc_ids": doc_ids, "module_name": module_name, "api_count": len(apis)}


def process_axure_zip(file_path: str, module_name: str = None, progress_cb=None) -> dict:
    """处理 Axure HTML 演示包：解析页面树 + UI 文本 + 交互 -> 存入 product_docs + SQLite。

    Args:
        file_path: Axure 导出的 .zip 文件路径
        module_name: 所属模块（如不指定则从 sitemap 自动提取）
        progress_cb: 可选，进度回调 (0~100, message)
    """
    from agent_components.axure_parser import AxureParser

    cb = progress_cb or (lambda p, m: None)
    logger.info(f"\n{'=' * 60}")
    logger.info(f"[Phase A] 处理 Axure 原型: {os.path.basename(file_path)}")

    db = get_chroma_db()
    file_name = os.path.basename(file_path)

    cb(5, "解压 Axure 包...")
    parser = AxureParser(file_path)
    try:
        cb(15, "解析页面树和 sitemap...")
        parsed = parser.parse()
        project_name = parsed.get("project_name", "Unknown")
        module = module_name or project_name
        cb(40, "提取 UI 文本和交互...")
        chunks = parser.to_product_doc_chunks(parsed)

        page_details = parsed.get("page_details", {})
        logger.info(f"   => 项目: {project_name}, 页面: {len(page_details)}")

        if not chunks:
            logger.warning(f"   ⚠️ Axure 解析后无内容（0 个页面），跳过入库")
            return {"doc_id": "", "module_name": module, "chunks": 0}

        # LLM 提取关联模块（复用 product_doc_extract_prompt 的语义分析能力）
        cb(70, "AI 分析关联模块...")
        graph = ChatTestAgentGraph()
        related = set()
        try:
            from prompts.extraction_prompts import product_doc_extract_prompt
            from prompts.response_model import DocModuleExtract
            prompt = product_doc_extract_prompt()
            # 取页面详情文本拼接，控制在单批上限内
            page_items = list(page_details.items())
            if len(page_items) > 50:
                logger.warning("Axure 页面数 %d > 50，已截断至 50 页用于 LLM 提取", len(page_items))
            detail_text = "\n".join(
                f"[{url}] {detail.get('ui_text', '')}"
                for url, detail in page_items[:50]
            )
            batch_limit = config.MAX_INGEST_CHARS_PER_BATCH
            if len(detail_text) > batch_limit:
                logger.warning("Axure 页面详情文本 %d 字符 > %d，已截断用于 LLM 提取",
                               len(detail_text), batch_limit)
                detail_text = detail_text[:batch_limit]  # Python str 切片，UTF-8 安全
            result = graph._invoke_structured(
                prompt, DocModuleExtract,
                method="json_mode",
                doc_text=detail_text,
            )
            related = set(result.related_modules or [])
            logger.info(f"   => LLM 识别关联模块: {related}")
        except Exception as e:
            logger.error("   => 关联模块分析失败: %s", e, exc_info=True)
        related.discard(module)

        # ── 提取每块 page_name（从 "## 页面: xxx" 行中解析）──
        import re as _re
        def _extract_page_name(chunk: str) -> str:
            m = _re.search(r'##\s*页面[：:]\s*(.+)', chunk)
            return m.group(1).strip() if m else ""

        # 5a. 写入 SQLite document_chunks 表（含 page_name）
        cb(75, "写入文档块...")
        doc_id = _safe_doc_id("axure", file_name, module)
        # to_product_doc_chunks 返回 list[dict]（{content, page_name}），统一解包为字符串
        # （2026-08-03 修复：原代码把 dict 当 str 传给 _extract_page_name → re 报
        #   "expected string or bytes-like object, got 'dict'"）
        chunk_texts = [
            c.get("content", c) if isinstance(c, dict) else c for c in chunks
        ]
        chunk_pages = [
            c.get("page_name", "") if isinstance(c, dict) else _extract_page_name(c)
            for c in chunks
        ]
        for i, content in enumerate(chunk_texts):
            _save_single_chunk(doc_id, i, content, page_name=chunk_pages[i])

        # 5b. 写入 documents 元数据
        _save_to_sqlite(
            doc_id=doc_id,
            file_name=file_name,
            file_type="zip",
            doc_type="axure",
            chunk_count=len(chunk_texts),
            module_name=module,
        )
        logger.info(f"   [SQLite] document_chunks + documents 入库完成 (doc_id={doc_id})")

        # 6. 批量生成 simple_summary（同步等待，失败写补偿任务）
        cb(80, "AI 生成摘要...")
        _generate_batch_summaries(doc_id, chunk_texts, file_name,
                                   progress_cb=cb,
                                   page_name=chunk_pages[0] if chunk_texts else "")

        # 7. 写入 ChromaDB（检索文本，失败时补偿回滚 SQLite）
        cb(90, "向量化入库中...")
        try:
            from database import get_session_ctx
            from database.models import DocumentChunk
            chunk_records = []
            with get_session_ctx() as session:
                chunk_records = session.query(DocumentChunk).filter_by(
                    doc_id=doc_id).order_by(DocumentChunk.chunk_index).all()
                chunk_records = [{
                    "content": c.content,
                    "simple_summary": c.simple_summary,
                    "page_name": c.page_name,
                } for c in chunk_records]
            search_texts = [_build_doc_search_text(c) for c in chunk_records]
            db.delete_by_doc_id(doc_id)
            db.add_product_doc_chunks(doc_id, search_texts, doc_type="axure")
            logger.info(f"   [ChromaDB] 入库完成 (doc_id={doc_id}), {len(chunks)} 块（检索文本）")
        except Exception:
            logger.error("   [ChromaDB] 写入失败，启动补偿回滚 SQLite", exc_info=True)
            _delete_sqlite_doc(doc_id)
            _delete_document_chunks(doc_id)
            raise

        cb(95, "入库完成")
        return {"doc_id": doc_id, "module_name": module, "chunks": len(chunks)}
    finally:
        parser.cleanup()


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="文档入库工具（CLI）")
    parser.add_argument("path", help="文档路径")
    parser.add_argument("--type", choices=["product", "api", "axure"], default="product",
                        help="文档类型（默认 product）")
    parser.add_argument("--module", default=None, help="所属模块名（可选）")
    args = parser.parse_args()

    if args.type == "api":
        result = process_api_doc(args.path, default_module=args.module)
    elif args.type == "axure":
        result = process_axure_zip(args.path, module_name=args.module)
    else:
        result = process_product_doc(args.path)

    logger.info(f"\n结果: {result}")
